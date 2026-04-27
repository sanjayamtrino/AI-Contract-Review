import io
import json
from typing import Any

from docx import Document
from fastapi import APIRouter, Depends, File, Form, Header, HTTPException, UploadFile

from src.api.session_utils import get_session_id
from src.dependencies import get_service_container
from src.schemas.contract_analyzer import ContractAnalyzerResponse
from src.schemas.describe_and_draft import DraftRequest, DraftResponse
from src.schemas.doc_chat import DocChatResponse
from src.schemas.general_review import GeneralReviewRequest, GeneralReviewResponse
from src.schemas.playbook_review import (
    PlayBookReviewFinalResponse,
    RuleCheckRequest,
    RuleInfo,
)
from src.tools.comparision import run as compare_documents_service
from src.tools.describe_and_draft import draft_document as draft_document_service
from src.tools.doc_chat import query_document as query_document_service
from src.tools.general_review import clause_review, full_document_review
from src.tools.key_information import (
    get_key_information_document as contract_analyzer_service,
)
from src.tools.new_playbook_review import (
    playbook_review_service as new_playbook_review_service,
)
from src.tools.playbook_review import review_document as playbook_review_service

router = APIRouter(tags=["agents"])


@router.post("/compare-documents")
async def compare_documents_endpoint(file_a: UploadFile, file_b: UploadFile, session_id: str = Header(..., alias="X-Session-Id")) -> Any:
    """Compare two documents and return their differences."""

    document_a = Document(io.BytesIO(await file_a.read()))
    document_b = Document(io.BytesIO(await file_b.read()))

    comparison_result = await compare_documents_service(session_id=session_id, document_a=document_a, document_b=document_b)
    return comparison_result


@router.post("/playbook-review", response_model=PlayBookReviewFinalResponse)
async def playbook_review_endpoint(request: RuleCheckRequest, session_id: str = Header(..., alias="X-Session-Id")) -> PlayBookReviewFinalResponse:
    """Run playbook validation checks."""

    review_result = await playbook_review_service(session_id=session_id, request=request)
    return review_result


@router.post("/new-playbook-review", response_model=PlayBookReviewFinalResponse)
async def new_playbook_review_endpoint(request: str = Form(...), document: UploadFile = File(...), session_id: str = Header(..., alias="X-Session-Id")):
    """Run playbook validation checks using the new clause extraction method."""

    request = RuleCheckRequest(**json.loads(request))

    document_content = Document(io.BytesIO(await document.read()))

    # Get session data from session manager
    service_container = get_service_container()
    session_data = service_container.session_manager.get_session(session_id)

    review_result = await new_playbook_review_service(request=request, document=document_content, session_data=session_data)

    return review_result


@router.post("/general-review", response_model=GeneralReviewResponse)
async def review_contract(request: GeneralReviewRequest, session_id: str = Depends(get_session_id)) -> GeneralReviewResponse:
    """Run the general review agent against an ingested document."""
    try:
        if request.selected_clause and request.selected_clause.strip():
            return await clause_review(
                session_id=session_id,
                clause_text=request.selected_clause,
                user_prompt=request.prompt,
                clause_title=(request.clause_title or "Selected Clause").strip() or "Selected Clause",
            )

        return await full_document_review(
            session_id=session_id,
            user_prompt=request.prompt,
        )

    except ValueError as err:
        raise HTTPException(status_code=400, detail=str(err))
    except Exception as err:
        raise HTTPException(status_code=500, detail=f"General review error: {str(err)}")


@router.post("/contract-analyzer")
async def contract_analyzer_endpoint(file: UploadFile, session_id: str = Header(..., alias="X-Session-Id")) -> ContractAnalyzerResponse:
    """Analyze a contract document and extract key information."""

    document = Document(io.BytesIO(await file.read()))
    document_data = "\n".join([para.text for para in document.paragraphs if para.text.strip() != ""])

    analysis_result: ContractAnalyzerResponse = await contract_analyzer_service(content=document_data, session_id=session_id)
    return analysis_result


@router.post("/query-document", response_model=DocChatResponse)
async def query_document_endpoint(query: str, session_id: str = Depends(get_session_id)) -> DocChatResponse:
    """Query the document chunks based on the given query and session ID."""

    llm_result = await query_document_service(query=query, session_id=session_id)
    return llm_result


@router.post("/draft", response_model=DraftResponse)
async def draft_document_endpoint(request: DraftRequest, session_id: str = Depends(get_session_id)) -> DraftResponse:
    """Draft the document/clause  for the user given query."""

    result = await draft_document_service(request, session_id)
    return result


# @router.post("/generate-nda-headings")
# async def generate_nda_endpoint(request: NDAGenerationHeadingRequest, session_id: str = Depends(get_session_id)) -> NDAGenerationHeadingResponse:
#     """Stepwise NDA generation endpoint."""

#     response = await generate_nda_headings(request, session_id)
#     return response


# @router.post("/generate-nda-content")
# async def generate_nda_content(request: NDAContentGenerationRequest, session_id: str = Depends(get_session_id)) -> NDAContentGenerationResponse:
#     """Generate content for a specific NDA heading based on type of document."""

#     response = await generate_heading_description(request, session_id)
#     return response
