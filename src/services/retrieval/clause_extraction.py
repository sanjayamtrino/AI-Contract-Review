# """
# extract_clauses.py
# ==================
# Extracts clauses from legal .docx files without using any LLM.

# Usage:
#     python extract_clauses.py <file.docx> [file2.docx ...]
#     python extract_clauses.py *.docx --output results.json
#     python extract_clauses.py *.docx --format json|yaml|text

# Output (default JSON):
#     [
#       {
#         "document": "NDA_v1.docx",
#         "clauses": [
#           {
#             "number": "1",
#             "title": "Definition of Confidential Information and Exclusions",
#             "content": "...",
#             "sub_clauses": [
#               {"number": "1(a)", "title": "", "content": "..."},
#               ...
#             ]
#           },
#           ...
#         ]
#       },
#       ...
#     ]
# """

# import argparse
# import json
# import re
# import sys
# from dataclasses import asdict, dataclass, field
# from pathlib import Path
# from typing import Optional

# from docx import Document

# # ---------------------------------------------------------------------------
# # Regex patterns that match typical legal clause numbering
# # ---------------------------------------------------------------------------

# # Top-level: "1.", "2.", "3A.", "4b." etc.
# TOP_LEVEL = re.compile(r"^(\d+[A-Za-z]?)\.\s+(.*)", re.DOTALL)

# # Sub-clause: "(a)", "(b)", "(i)", "(ii)", "(iii)" etc.
# SUB_LEVEL = re.compile(r"^\(([a-z]+|[ivxlcdm]+|\d+)\)\s+(.*)", re.DOTALL)

# # Roman numeral sub-sub: "(i)", "(ii)", …
# ROMAN = re.compile(r"^[ivxlcdm]+$")

# # Patterns that look like clause headers but are actually metadata / footers
# SKIP_PATTERNS = [
#     re.compile(r"^\s*$"),  # blank
#     re.compile(r"Initials\s*:", re.I),  # signature lines
#     re.compile(r"^(IN WITNESS WHEREOF|WHEREAS|NOW,? THERE)", re.I),
#     re.compile(r"^RECITALS", re.I),
#     re.compile(r"Rev \d{2}-\d{2}-\d{4}"),  # revision stamps
#     re.compile(r"^\s*-\s*\d+\s*-\s*"),  # page number lines
# ]

# # Styles in the MSA that carry explicit level information
# LEGAL2_L1_STYLE = "Legal2_L1"
# LEGAL2_L2_STYLE = "Legal2_L2"


# # ---------------------------------------------------------------------------
# # Data classes
# # ---------------------------------------------------------------------------


# @dataclass
# class SubClause:
#     number: str
#     title: str
#     content: str


# @dataclass
# class Clause:
#     number: str
#     title: str
#     content: str
#     sub_clauses: list = field(default_factory=list)


# @dataclass
# class DocumentResult:
#     document: str
#     clauses: list = field(default_factory=list)


# # ---------------------------------------------------------------------------
# # Helpers
# # ---------------------------------------------------------------------------


# def _should_skip(text: str) -> bool:
#     for pat in SKIP_PATTERNS:
#         if pat.search(text):
#             return True
#     return False


# def _split_title_and_content(raw: str):
#     """
#     Given raw text like "Definition of Confidential Information.  Body text…"
#     split into (title, body).  Title ends at the first sentence-ending period
#     followed by two spaces, a newline, or end-of-string.
#     """
#     # Try splitting on ". " with a capital letter following (likely body)
#     m = re.search(r"\.\s{2,}", raw)
#     if m:
#         return raw[: m.start() + 1].strip(), raw[m.end() :].strip()
#     # Try splitting after first period if the rest has content
#     m = re.search(r"\.\s+([A-Z])", raw)
#     if m:
#         return raw[: m.start() + 1].strip(), raw[m.start() + 1 :].strip()
#     return raw.strip(), ""


# def _clean(text: str) -> str:
#     """Collapse excessive whitespace."""
#     text = re.sub(r"[ \t]+", " ", text)
#     text = re.sub(r"\n{3,}", "\n\n", text)
#     return text.strip()


# # ---------------------------------------------------------------------------
# # Per-document extraction strategies
# # ---------------------------------------------------------------------------


# def _extract_plain_numbered(doc: Document) -> list:
#     """
#     Strategy for docs where clauses are plain 'Normal' paragraphs starting
#     with a number pattern: NDA and Subcontractor Agreement.
#     """
#     clauses: list[Clause] = []
#     current_clause: Optional[Clause] = None
#     current_sub: Optional[SubClause] = None

#     for para in doc.paragraphs:
#         text = para.text.strip()

#         if not text or _should_skip(text):
#             continue

#         m_top = TOP_LEVEL.match(text)
#         m_sub = SUB_LEVEL.match(text)

#         if m_top:
#             # Finalise any open sub-clause
#             if current_sub is not None and current_clause is not None:
#                 current_clause.sub_clauses.append(current_sub)
#                 current_sub = None

#             num = m_top.group(1)
#             rest = m_top.group(2).strip()
#             title, body = _split_title_and_content(rest)

#             current_clause = Clause(
#                 number=num,
#                 title=_clean(title),
#                 content=_clean(body),
#             )
#             clauses.append(current_clause)

#         elif m_sub and current_clause is not None:
#             # Finalise previous sub-clause
#             if current_sub is not None:
#                 current_clause.sub_clauses.append(current_sub)

#             sub_num = m_sub.group(1)
#             sub_rest = m_sub.group(2).strip()
#             sub_title, sub_body = _split_title_and_content(sub_rest)

#             full_num = f"{current_clause.number}({sub_num})"
#             current_sub = SubClause(
#                 number=full_num,
#                 title=_clean(sub_title) if sub_title != sub_body else "",
#                 content=_clean(sub_body if sub_body else sub_title),
#             )

#         else:
#             # Continuation paragraph — append to the right bucket
#             if current_sub is not None:
#                 current_sub.content = _clean(current_sub.content + " " + text)
#             elif current_clause is not None:
#                 current_clause.content = _clean(current_clause.content + " " + text)

#     # Flush last open sub-clause
#     if current_sub is not None and current_clause is not None:
#         current_clause.sub_clauses.append(current_sub)

#     return clauses


# def _extract_styled(doc: Document) -> list:
#     """
#     Strategy for docs using explicit heading styles (Legal2_L1 / Legal2_L2 or
#     Heading 1 / Heading 2): Master Services Agreement, EPIT Agreement.
#     """
#     clauses: list[Clause] = []
#     current_clause: Optional[Clause] = None
#     current_sub: Optional[SubClause] = None
#     clause_counter = 0
#     sub_counter = 0

#     # Detect auto-numbering from numId attribute
#     WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

#     def has_autonumber(para):
#         try:
#             return para._element.find(f"{{{WNS}}}pPr/{{{WNS}}}numPr/{{{WNS}}}numId") is not None
#         except Exception:
#             return False

#     for para in doc.paragraphs:
#         text = para.text.strip()
#         style = para.style.name

#         if not text or _should_skip(text):
#             continue

#         is_l1 = style == LEGAL2_L1_STYLE or style.startswith("Heading 1") or (has_autonumber(para) and style in ("Normal", "List Paragraph") and not SUB_LEVEL.match(text))

#         is_l2 = style == LEGAL2_L2_STYLE or style.startswith("Heading 2")

#         # Also detect plain-numbered top-level in styled docs
#         m_top = TOP_LEVEL.match(text)
#         m_sub = SUB_LEVEL.match(text)

#         if m_top and not is_l2:
#             # e.g., "1. Services:" in Subcontractor
#             if current_sub is not None and current_clause is not None:
#                 current_clause.sub_clauses.append(current_sub)
#                 current_sub = None
#             num = m_top.group(1)
#             rest = m_top.group(2).strip()
#             title, body = _split_title_and_content(rest)
#             current_clause = Clause(number=num, title=_clean(title), content=_clean(body))
#             clauses.append(current_clause)

#         elif is_l1 and not m_top and not m_sub:
#             # Styled heading = new top-level clause
#             if current_sub is not None and current_clause is not None:
#                 current_clause.sub_clauses.append(current_sub)
#                 current_sub = None
#             clause_counter += 1
#             sub_counter = 0
#             title, body = _split_title_and_content(text)
#             current_clause = Clause(
#                 number=str(clause_counter),
#                 title=_clean(title),
#                 content=_clean(body),
#             )
#             clauses.append(current_clause)

#         elif is_l2 and current_clause is not None:
#             if current_sub is not None:
#                 current_clause.sub_clauses.append(current_sub)
#             sub_counter += 1
#             title, body = _split_title_and_content(text)
#             current_sub = SubClause(
#                 number=f"{current_clause.number}.{sub_counter}",
#                 title=_clean(title),
#                 content=_clean(body),
#             )

#         elif m_sub and current_clause is not None:
#             if current_sub is not None:
#                 current_clause.sub_clauses.append(current_sub)
#             sub_num = m_sub.group(1)
#             sub_rest = m_sub.group(2).strip()
#             sub_title, sub_body = _split_title_and_content(sub_rest)
#             current_sub = SubClause(
#                 number=f"{current_clause.number}({sub_num})",
#                 title=_clean(sub_title) if sub_title != sub_body else "",
#                 content=_clean(sub_body if sub_body else sub_title),
#             )

#         else:
#             # Continuation
#             if current_sub is not None:
#                 current_sub.content = _clean(current_sub.content + " " + text)
#             elif current_clause is not None:
#                 current_clause.content = _clean(current_clause.content + " " + text)

#     if current_sub is not None and current_clause is not None:
#         current_clause.sub_clauses.append(current_sub)

#     return clauses


# def _detect_strategy(doc: Document) -> str:
#     """
#     Heuristically decide which extraction strategy to use based on the styles
#     present in the document.
#     """
#     style_names = {p.style.name for p in doc.paragraphs if p.text.strip()}
#     if LEGAL2_L1_STYLE in style_names or "Heading 1" in style_names:
#         return "styled"
#     return "plain"


# # ---------------------------------------------------------------------------
# # Public API
# # ---------------------------------------------------------------------------


# def extract_clauses(docx_path: str) -> DocumentResult:
#     """
#     Extract all clauses from a .docx legal document.
#     Returns a DocumentResult with the file name and list of Clause objects.
#     """
#     doc = Document(docx_path)
#     strategy = _detect_strategy(doc)

#     if strategy == "styled":
#         clauses = _extract_styled(doc)
#     else:
#         clauses = _extract_plain_numbered(doc)

#     return DocumentResult(
#         document=Path(docx_path).name,
#         clauses=clauses,
#     )


# def result_to_dict(result: DocumentResult) -> dict:
#     """Convert a DocumentResult to a plain dict (JSON-serialisable)."""
#     out = {"document": result.document, "clauses": []}
#     for c in result.clauses:
#         cd = {
#             "number": c.number,
#             "title": c.title,
#             "content": c.content,
#         }
#         if c.sub_clauses:
#             cd["sub_clauses"] = [{"number": s.number, "title": s.title, "content": s.content} for s in c.sub_clauses]
#         out["clauses"].append(cd)
#     return out


# def format_text(result: DocumentResult) -> str:
#     """Human-readable plain-text output."""
#     lines = [f"{'='*70}", f"DOCUMENT: {result.document}", f"{'='*70}"]
#     for c in result.clauses:
#         lines.append(f"\n[Clause {c.number}] {c.title}")
#         lines.append("-" * 60)
#         if c.content:
#             lines.append(c.content)
#         for s in c.sub_clauses:
#             sub_label = f"  [{s.number}]"
#             if s.title:
#                 sub_label += f" {s.title}"
#             lines.append(sub_label)
#             if s.content:
#                 lines.append(f"    {s.content}")
#     return "\n".join(lines)


# # ---------------------------------------------------------------------------
# # CLI
# # ---------------------------------------------------------------------------


# def main():
#     parser = argparse.ArgumentParser(description="Extract clauses from legal .docx files (no LLM required).")
#     parser.add_argument("files", nargs="+", help=".docx file(s) to process")
#     parser.add_argument("--format", choices=["json", "text"], default="json", help="Output format (default: json)")
#     parser.add_argument("--output", "-o", default=None, help="Write output to this file (default: stdout)")
#     args = parser.parse_args()

#     all_results = []
#     for fp in args.files:
#         try:
#             result = extract_clauses(fp)
#             all_results.append(result)
#             print(f"✓ Processed: {fp}  ({len(result.clauses)} top-level clauses found)", file=sys.stderr)
#         except Exception as e:
#             print(f"✗ Error processing {fp}: {e}", file=sys.stderr)

#     if args.format == "json":
#         output = json.dumps([result_to_dict(r) for r in all_results], indent=2, ensure_ascii=False)
#     else:
#         output = "\n\n".join(format_text(r) for r in all_results)

#     if args.output:
#         Path(args.output).write_text(output, encoding="utf-8")
#         print(f"\nOutput written to: {args.output}", file=sys.stderr)
#     else:
#         print(output)


# if __name__ == "__main__":
#     main()

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from docx.document import Document

# ---------------------------------------------------------------------------
# Constants & patterns
# ---------------------------------------------------------------------------

TOP_LEVEL = re.compile(r"^(\d+[A-Za-z]{0,2})\.\s*(.*)", re.DOTALL)
SUB_LEVEL = re.compile(r"^\(([a-z]+|[ivxlcdm]+|\d+)\)\s+(.*)", re.DOTALL)

SKIP_PATTERNS = [
    re.compile(r"^\s*$"),
    re.compile(r"Initials\s*:", re.I),
    re.compile(r"^(IN WITNESS WHEREOF|WHEREAS|NOW,?\s*THERE)", re.I),
    re.compile(r"^RECITALS", re.I),
    re.compile(r"Rev \d{2}-\d{2}-\d{4}"),
    re.compile(r"^\s*-\s*\d+\s*-\s*"),
    re.compile(r"^(Contractor|Subcontractor):\s*$", re.I),
    re.compile(r"^(BY|NAME|TITLE|DATE):\s*$", re.I),
    re.compile(r"^\(Authorized Signature\)$", re.I),
    re.compile(r"^\(Print Name", re.I),
    re.compile(r"^\(Title of Signatory\)$", re.I),
    re.compile(r"^\(Execution Date\)", re.I),
]

LEGAL2_L1 = "Legal2_L1"
LEGAL2_L2 = "Legal2_L2"

# A "label" is at most this many words
# Word limit for colon-separated labels  "Independent Company: body…"
COLON_LABEL_MAX = 7
# Word limit for double-space-separated labels  "Return of Advances.  body…"
# Higher because double-space is a deliberate, unambiguous separator in legal docs
PERIOD_LABEL_MAX = 11


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------


@dataclass
class SubClause:
    number: str
    title: str
    content: str


@dataclass
class Clause:
    number: str
    title: str
    content: str
    sub_clauses: list = field(default_factory=list)


@dataclass
class DocumentResult:
    document: str
    clauses: list = field(default_factory=list)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _should_skip(text: str) -> bool:
    return any(p.search(text) for p in SKIP_PATTERNS)


def _is_label(candidate: str, max_words: int = COLON_LABEL_MAX) -> bool:
    """Return True if candidate looks like a short clause label (≤ max_words words)."""
    return 1 <= len(candidate.strip().split()) <= max_words


def _split(raw: str) -> tuple:
    """
    Split raw clause text into (title, content).

    Rules in priority order:
      1. Trailing colon only  e.g. "Termination:"
         → title = stripped label, content = ""
      2. Trailing period only  e.g. "Miscellaneous."
         → title = stripped label, content = ""
      3. Colon + body   e.g. "Independent Company: Subcontractor and…"
         → title = left (if ≤ LABEL_MAX_WORDS words), content = right
      4. Period + body  e.g. "Remedies. Money damages may not…"
         → title = left (if ≤ LABEL_MAX_WORDS words), content = right
      5. Fallback: pure body paragraph
         → title = "", content = whole text
    """
    raw = raw.strip()
    if not raw:
        return "", ""

    # 1. Whole string ends with colon and is a short label
    if raw.endswith(":") and _is_label(raw.rstrip(":")):
        return raw.rstrip(":").strip(), ""

    # 2. Whole string ends with period and is a short label
    if raw.endswith(".") and _is_label(raw.rstrip(".")):
        return raw.rstrip(".").strip(), ""

    # 3. Colon split — "Short Label: body text…"
    m = re.search(r":\s+", raw)
    if m:
        left = raw[: m.start()].strip()
        right = raw[m.end() :].strip()
        if _is_label(left, COLON_LABEL_MAX) and right:
            return left, right

    # 4. Period + 2+ spaces — "Short Label.  Body text…"  (deliberate separator, allow longer labels)
    m = re.search(r"\.\s{2,}", raw)
    if m:
        left = raw[: m.start()].strip()
        right = raw[m.end() :].strip()
        if _is_label(left, PERIOD_LABEL_MAX) and right:
            return left, right

    # 4b. Period + single space — "Short Label. Body text…"  (only for very short labels)
    m = re.search(r"\.\s+", raw)
    if m:
        left = raw[: m.start()].strip()
        right = raw[m.end() :].strip()
        if _is_label(left, PERIOD_LABEL_MAX) and right:
            return left, right

    # 5. Fallback — body paragraph with no title
    return "", raw


def _clean(text: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", re.sub(r"[ \t]+", " ", text)).strip()


def _dedup(clauses: list) -> list:
    """Append b/c/d… to duplicate clause numbers."""
    seen: dict = {}
    suffixes = "bcdefghijklmnopqrstuvwxyz"
    for c in clauses:
        key = c.number
        if key in seen:
            idx = seen[key]
            seen[key] += 1
            c.number = f"{key}{suffixes[idx]}"
        else:
            seen[key] = 0
    return clauses


# ---------------------------------------------------------------------------
# Extraction strategies
# ---------------------------------------------------------------------------


def _extract_plain(doc: Document) -> list:
    """For docs with plain numbered paragraphs (NDA, Subcontractor)."""
    clauses: list[Clause] = []
    cur: Optional[Clause] = None
    sub: Optional[SubClause] = None

    def flush():
        nonlocal sub
        if sub and cur:
            cur.sub_clauses.append(sub)
            sub = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or _should_skip(text):
            continue

        mt = TOP_LEVEL.match(text)
        ms = SUB_LEVEL.match(text)

        if mt:
            flush()
            num = mt.group(1)
            rest = mt.group(2).strip()
            t, b = _split(rest)
            cur = Clause(number=num, title=_clean(t), content=_clean(b))
            clauses.append(cur)

        elif ms and cur:
            flush()
            snum = ms.group(1)
            srest = ms.group(2).strip()
            st, sb = _split(srest)
            sub = SubClause(
                number=f"{cur.number}({snum})",
                title=_clean(st),
                content=_clean(sb),
            )

        else:
            if sub:
                sub.content = _clean(sub.content + " " + text)
            elif cur:
                cur.content = _clean(cur.content + " " + text)

    flush()
    return _dedup(clauses)


def _extract_styled(doc: Document) -> list:
    """For docs using heading styles (MSA, EPIT)."""
    clauses: list[Clause] = []
    cur: Optional[Clause] = None
    sub: Optional[SubClause] = None
    cidx = 0
    sidx = 0

    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def autonumbered(para):
        try:
            return para._element.find(f"{{{WNS}}}pPr/{{{WNS}}}numPr/{{{WNS}}}numId") is not None
        except Exception:
            return False

    def flush():
        nonlocal sub
        if sub and cur:
            cur.sub_clauses.append(sub)
            sub = None

    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name
        if not text or _should_skip(text):
            continue

        mt = TOP_LEVEL.match(text)
        ms = SUB_LEVEL.match(text)

        is_l1 = style == LEGAL2_L1 or style.startswith("Heading 1") or (autonumbered(para) and style in ("Normal", "List Paragraph") and not ms)
        is_l2 = style == LEGAL2_L2 or style.startswith("Heading 2")

        if mt and not is_l2:
            flush()
            num = mt.group(1)
            rest = mt.group(2).strip()
            t, b = _split(rest)
            cur = Clause(number=num, title=_clean(t), content=_clean(b))
            clauses.append(cur)

        elif is_l1 and not mt and not ms:
            flush()
            cidx += 1
            sidx = 0
            t, b = _split(text)
            cur = Clause(number=str(cidx), title=_clean(t), content=_clean(b))
            clauses.append(cur)

        elif is_l2 and cur:
            flush()
            sidx += 1
            t, b = _split(text)
            sub = SubClause(
                number=f"{cur.number}.{sidx}",
                title=_clean(t),
                content=_clean(b),
            )

        elif ms and cur:
            flush()
            snum = ms.group(1)
            srest = ms.group(2).strip()
            st, sb = _split(srest)
            sub = SubClause(
                number=f"{cur.number}({snum})",
                title=_clean(st),
                content=_clean(sb),
            )

        else:
            if sub:
                sub.content = _clean(sub.content + " " + text)
            elif cur:
                cur.content = _clean(cur.content + " " + text)

    flush()
    return _dedup(clauses)


def _detect(doc: Document) -> str:
    styles = {p.style.name for p in doc.paragraphs if p.text.strip()}
    return "styled" if (LEGAL2_L1 in styles or "Heading 1" in styles) else "plain"


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def extract_clauses(docx_path: str) -> DocumentResult:
    doc = Document(docx_path)
    fn = _extract_styled if _detect(doc) == "styled" else _extract_plain
    return DocumentResult(document=Path(docx_path).name, clauses=fn(doc))


def result_to_dict(r: DocumentResult) -> dict:
    out: dict = {"document": r.document, "clauses": []}
    for c in r.clauses:
        cd: dict = {"number": c.number, "title": c.title, "content": c.content}
        if c.sub_clauses:
            cd["sub_clauses"] = [{"number": s.number, "title": s.title, "content": s.content} for s in c.sub_clauses]
        out["clauses"].append(cd)
    return out


def format_text(r: DocumentResult) -> str:
    lines = [f"{'='*70}", f"DOCUMENT: {r.document}", f"{'='*70}"]
    for c in r.clauses:
        lines.append(f"\n[Clause {c.number}] {c.title or '(no title)'}")
        lines.append("-" * 60)
        if c.content:
            lines.append(c.content)
        for s in c.sub_clauses:
            lbl = f"  [{s.number}]" + (f" {s.title}" if s.title else "")
            lines.append(lbl)
            if s.content:
                lines.append(f"    {s.content}")
    return "\n".join(lines)
