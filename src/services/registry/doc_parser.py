import re
import time
import uuid
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document as DocxDocument
from docx.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from src.config.logging import Logger
from src.config.settings import get_settings
from src.exceptions.parser_exceptions import (
    DocxCleaningException,
    DocxMetadataExtractionException,
    DocxParagraphExtractionException,
    DocxTableExtractionException,
)
from src.schemas.registry import Chunk, ParseResult
from src.services.registry.base_parser import BaseParser
from src.services.session_manager import SessionData
from src.services.vector_store.embeddings.embedding_service import (
    BGEEmbeddingService,
    # HuggingFaceEmbeddingService,
)

# from src.services.vector_store.embeddings.jina_embeddings import JinaEmbeddings
# from src.services.vector_store.embeddings.gemini_embeddings import (
#     GeminiEmbeddingService,
# )
# from src.services.vector_store.embeddings.openai_embeddings import OpenAIEmbeddings
from src.services.vector_store.manager import get_faiss_vector_store, index_chunks


class DocxParser(BaseParser, Logger):
    """Parser for DOCX documents."""

    def __init__(self) -> None:
        """Initialize the DOCX parser."""

        super().__init__()
        self.settings = get_settings()
        # self.embedding_service = HuggingFaceEmbeddingService()
        self.embedding_service = BGEEmbeddingService()
        # self.embedding_service = GeminiEmbeddingService()
        # self.embedding_service = OpenAIEmbeddings()
        # self.embedding_service = JinaEmbeddings()
        self.vector_store = get_faiss_vector_store(self.embedding_service.get_embedding_dimensions())

    def _clean_text(self, text: str) -> str:
        """Cleans and normalize the text content."""

        if not text:
            return ""

        # Remove excessive whitespace and newlines (collapse all whitespace to single space)
        text = re.sub(r"\s+", " ", text)

        # Remove leading/trailing whitespace
        text = text.strip()

        # Remove multiple spaces (if any remain)
        text = re.sub(r" {2,}", " ", text)

        # Normalize common special characters
        text = text.replace("\u00a0", " ")  # Non-breaking space
        text = text.replace("\u200b", "")  # Zero-width space
        text = text.replace("\ufeff", "")  # Zero-width no-break space
        text = text.replace("\r", "")  # Carriage return

        # Remove any remaining control characters except normal space
        text = re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]", "", text)
        text = text.lstrip(" .\n\t")

        return text

    async def clean_document(self, document: Document) -> None:
        """Clean the document before parsing to remove unwanted elements."""

        try:
            for paragraph in document.paragraphs:
                if paragraph.text:
                    paragraph.text = " ".join(paragraph.text.split())

            self.logger.debug("Document cleaned successfully.")
        except Exception as e:
            self.logger.error(f"Error cleaning document: {e}")
            raise DocxCleaningException(f"Error cleaning document: {e}") from e

    async def _extract_metadata(self, document: Document) -> Dict[str, Any]:
        """Extract metadata about the document."""

        try:
            properties = document.core_properties

            # Count words from paragraphs
            paragraph_word_count = sum(len(p.text.split()) for p in document.paragraphs)

            # Count words from tables
            table_word_count = 0
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        table_word_count += len(cell.text.split())

            # Total word count including both paragraphs and tables
            total_word_count = paragraph_word_count + table_word_count

            # assign a unique document id for this parse
            document_id = str(uuid.uuid4())

            metadata: Dict[str, Any] = {
                "source": "docx",
                "document_id": document_id,
                "author": properties.author or "Unknown",
                "title": properties.title or "Untitled",
                "subject": properties.subject or "",
                "created_at": properties.created.isoformat() if properties.created else None,
                "modified_at": properties.modified.isoformat() if properties.modified else None,
                "last_modified_by": properties.last_modified_by or "Unknown",
                "revision": properties.revision or 0,
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "image_count": len(document.inline_shapes),
                "word_count": total_word_count,
                "paragraph_word_count": paragraph_word_count,
                "table_word_count": table_word_count,
            }

            return metadata
        except Exception as e:
            self.logger.error(f"Error extracting metadata: {e}")
            raise DocxMetadataExtractionException(f"Error extracting metadata: {e}") from e

    async def _extract_paragraphs(self, document: Document) -> List[Dict[str, Any]]:
        """Extract paragraphs from the document."""

        paragraphs_data: List[Dict[str, Any]] = []

        try:
            for index, paragraph in enumerate(document.paragraphs):
                if paragraph.text.strip():
                    cleaned_text = self._clean_text(paragraph.text)
                    if cleaned_text:
                        paragraphs_data.append(
                            {
                                "index": index,
                                "content": cleaned_text,
                                "is_heading": paragraph.style.name.startswith("Heading") if paragraph.style else False,
                            }
                        )

        except Exception as e:
            self.logger.error(f"Error extracting paragraphs: {e}")
            raise DocxParagraphExtractionException(f"Error extracting paragraphs: {e}") from e

        return paragraphs_data

    async def _extract_images(self, document: Document) -> List[Dict[str, Any]]:
        """Extract images from the document."""

        return []

    async def _extract_tables(self, document: Document) -> List[Dict[str, Any]]:
        """Extract tables from the document."""

        tables_data: List[Dict[str, Any]] = []

        try:
            for table_index, table in enumerate(document.tables):
                table_content = []
                for row in table.rows:
                    row_data = [self._clean_text(cell.text) for cell in row.cells]
                    table_content.append(row_data)

                tables_data.append(
                    {
                        "table_index": table_index,
                        "content": table_content,
                    }
                )

        except Exception as e:
            self.logger.error(f"Error extracting tables: {e}")
            raise DocxTableExtractionException(f"Error extracting tables: {e}") from e

        return tables_data

    async def _get_text_splitter(self) -> RecursiveCharacterTextSplitter:
        """Get the text splitter for chunking the document content."""

        return RecursiveCharacterTextSplitter(
            chunk_size=self.settings.chunk_size,
            chunk_overlap=self.settings.chunk_overlap,
            separators=[
                "\n\n\n",  # Multiple newlines (section breaks)
                "\n\n",  # Paragraph breaks
                "\n",  # Line breaks
                ". ",  # Sentence ends
                "! ",  # Exclamation
                "? ",  # Question
                "; ",  # Semicolon
                ": ",  # Colon
                ", ",  # Comma
                " ",  # Space
                "",  # Character level (last resort)
            ],
            length_function=len,
            keep_separator=True,
            is_separator_regex=False,
        )

    async def parse(self, document: Document, session_data: Optional["SessionData"] = None) -> ParseResult:
        """Parse the DOCX data."""

        start_time = time.time()

        try:
            self.logger.info("Starting DOCX parsing.")

            # Clean the document
            await self.clean_document(document=document)

            # Extract metadata
            metadata = await self._extract_metadata(document=document)
            document_id = metadata.get("document_id")

            # Determine which vector store to use
            vector_store = session_data.vector_store if session_data else self.vector_store

            # Extract paragraphs
            paragraphs = await self._extract_paragraphs(document=document)

            # Extract tables
            tables = await self._extract_tables(document=document)

            full_text = " ".join([p["content"] for p in paragraphs])
            full_text = self._clean_text(full_text)

            text_splitter = await self._get_text_splitter()
            chunks: List[Chunk] = []
            chunk_index = 0

            # Create chunks from paragraph text
            if full_text:
                texts: List[str] = text_splitter.split_text(full_text)

                for text in texts:
                    cleaned_chunk = self._clean_text(text)

                    if not cleaned_chunk:
                        continue

                    self.logger.debug(f"Paragraph chunk {chunk_index} created with length {len(cleaned_chunk)}.")

                    # Embed the text
                    vector_data: List[float] = await self.embedding_service.generate_embeddings(text=cleaned_chunk, task="text-matching")
                    await vector_store.index_embedding(embedding=vector_data)

                    chunk = Chunk(
                        chunk_id=str(uuid.uuid4()),
                        document_id=document_id,
                        chunk_index=chunk_index,
                        content=cleaned_chunk,
                        embedding_model=self.embedding_service.model_name,
                        embedding_vector=vector_data,
                        metadata={
                            "chunk_type": "paragraph",
                        },
                        created_at=datetime.utcnow().isoformat(),
                    )
                    chunks.append(chunk)
                    chunk_index += 1

            # Create separate chunks for each table
            for table_data in tables:
                # Convert table to text format
                table_text_rows = []
                for row in table_data["content"]:
                    # Joining cells with pipe separator for readability
                    row_text = " | ".join(row)
                    table_text_rows.append(row_text)

                # Join all rows
                table_text = " ".join(table_text_rows)

                # Clean the table text
                cleaned_table_text = self._clean_text(table_text)

                if not cleaned_table_text:  # Skip empty tables
                    continue

                self.logger.debug(f"Table chunk {chunk_index} created with length {len(cleaned_table_text)}.")

                # Embed the table text
                vector_data: List[float] = await self.embedding_service.generate_embeddings(text=cleaned_table_text)
                await vector_store.index_embedding(embedding=vector_data)

                chunk = Chunk(
                    chunk_id=str(uuid.uuid4()),
                    document_id=document_id,
                    chunk_index=chunk_index,
                    content=cleaned_table_text,
                    embedding_model=self.embedding_service.model_name,
                    embedding_vector=vector_data,
                    metadata={
                        "chunk_type": "table",
                        "table_index": table_data["table_index"],
                        "row_count": len(table_data["content"]),
                        "column_count": len(table_data["content"][0]) if table_data["content"] else 0,
                    },
                    created_at=datetime.utcnow().isoformat(),
                )
                chunks.append(chunk)
                chunk_index += 1

            # Store chunks in the shared manager so RetrievalService can access them
            index_chunks(chunks)

            processing_time = time.time() - start_time

            return ParseResult(
                success=True,
                chunks=chunks,
                metadata=metadata,
                error_message=None,
                processing_time=processing_time,
            )
        except Exception as e:
            self.logger.error(f"Error parsing DOCX document: {e}")
            return ParseResult(
                success=False,
                chunks=[],
                metadata={},
                error_message=str(e),
                processing_time=0.0,
            )

    async def _index_embeddings(self, chunks: List[Chunk]) -> None:
        """Index the document embeddings into the vector store."""

        try:
            # start_time = time.time()

            for chunk in chunks:
                self.vector_store.add(chunk.embedding_vector)

        except Exception as e:
            raise ValueError("Unable to index the document into the vector store.") from e

    async def _get_health_status(self) -> Dict[str, Any]:
        """Get the health status of the DOCX parser."""

        status: str = "healthy"
        info: Dict[str, Any] = {}

        # Settings accessibility check
        try:
            _ = self.settings.chunk_size
            _ = self.settings.chunk_overlap
            info["settings_accessible"] = True
        except Exception as e:
            status = "unhealthy"
            info["settings_accessible"] = False
            info["error"] = str(e)

        # Text splitter check
        try:
            _ = await self._get_text_splitter()
            info["text_splitter_accessible"] = True
        except Exception as e:
            status = "unhealthy"
            info["text_splitter_accessible"] = False
            info["error"] = str(e)

        # DOCX test parsing check
        try:

            sample_docx = BytesIO()
            doc = DocxDocument()
            doc.add_paragraph("This is a test paragraph.")
            doc.save(sample_docx)
            sample_docx.seek(0)

            document = DocxDocument(sample_docx)
            parse_result = await self.parse(document=document)

            if parse_result.success:
                info["docx_parsing"] = "successful"
            else:
                status = "unhealthy"
                info["docx_parsing"] = "failed"
                info["error"] = parse_result.error_message
        except Exception as e:
            status = "unhealthy"
            info["docx_parsing"] = "failed"
            info["error"] = str(e)

        return {
            "status": status,
            "info": info,
        }

    async def is_healthy(self) -> bool:
        """Checks if the Parser is in healthy condition."""

        status: Dict[str, Any] = await self._get_health_status()
        return True if status.get("status", "") == "healthy" else False
