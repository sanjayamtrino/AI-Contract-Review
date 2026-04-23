# import re
# import time
# import uuid
# from datetime import datetime
# from typing import Any, Dict, List, Optional

# import numpy as np
# from docx.document import Document
# from numpy import dot
# from numpy.linalg import norm

# from src.config.logging import Logger
# from src.config.settings import get_settings
# from src.exceptions.parser_exceptions import (
#     DocxCleaningException,
#     DocxMetadataExtractionException,
#     DocxParagraphExtractionException,
#     DocxTableExtractionException,
#     EmptyTextException,
# )
# from src.schemas.playbook_review import TextInfo
# from src.schemas.registry import Chunk, ParseResult
# from src.services.registry.base_parser import BaseParser
# from src.services.session_manager import SessionData
# from src.services.vector_store.embeddings.base_embedding_service import (
#     BaseEmbeddingService,
# )
# from src.services.vector_store.manager import get_faiss_vector_store

# _SECTION_LABEL_RE = re.compile(
#     r"^("
#     r"\d+[\.\)]?\s+\S.*|"
#     r"\d+[\.\)]?\s*$|"  # "1."  /  "2)"
#     r"[A-Z][A-Z\s\.\,\&\'\-]{1,60}$"
#     r")"
# )

# # Regex: multi-word inline clause headings like "Audit Rights. " or "No Warranty. "
# _INLINE_CLAUSE_HEADING_RE = re.compile(
#     r"(?:^|(?<=\. ))" r"((?:[A-Z][a-z]+(?:'[a-z]+)?)" r"(?:" r"(?:\s+(?:and|&|of|the|to|for|on|in|or|with|at|by|an|a|your|its|from))*" r"\s+[A-Z][a-z]+(?:'[a-z]+)?" r")+" r"\.)\s"
# )

# # Regex: single-word clause titles like "Term. ", "Assignment. ", "Definitions. ".
# # Constrained to first word with 4+ total characters to exclude "Mr", "Dr", "Ms",
# # etc., and must be followed by a capital letter so we only match clause-body
# # starts, not mid-sentence capitalization.
# _SINGLE_WORD_HEADING_RE = re.compile(r"(?:^|(?<=\. ))" r"([A-Z][a-z]{3,}(?:'[a-z]+)?\.)" r"\s(?=[A-Z])")

# # Words that look like single-word clause titles but are actually sentence
# # adverbs or conjunctions. Anything in this set is ignored when matched by
# # _SINGLE_WORD_HEADING_RE.
# _SINGLE_WORD_HEADING_BLACKLIST = frozenset(
#     {
#         "however",
#         "further",
#         "furthermore",
#         "additionally",
#         "moreover",
#         "otherwise",
#         "nevertheless",
#         "notwithstanding",
#         "accordingly",
#         "therefore",
#         "thus",
#         "hence",
#         "thereby",
#         "thereafter",
#         "whereas",
#         "specifically",
#         "particularly",
#         "collectively",
#         "individually",
#         "alternatively",
#         "consequently",
#         "subsequently",
#         "respectively",
#         "conversely",
#         "similarly",
#         "meanwhile",
#         "indeed",
#         "finally",
#         "firstly",
#         "secondly",
#         "thirdly",
#         "lastly",
#         "instead",
#     }
# )


# def _find_clause_heading_matches(text: str) -> List[Dict[str, Any]]:
#     """Return clause-heading matches in paragraph order, deduped by start position.

#     Each entry is {"start": int, "end": int, "heading": str} where end points
#     to the first character AFTER the heading (i.e. start of body text).
#     """
#     found: Dict[int, Dict[str, Any]] = {}

#     for m in _INLINE_CLAUSE_HEADING_RE.finditer(text):
#         found[m.start()] = {"start": m.start(), "end": m.end(1) + 1, "heading": m.group(1)}

#     for m in _SINGLE_WORD_HEADING_RE.finditer(text):
#         word = m.group(1).rstrip(".").lower()
#         if word in _SINGLE_WORD_HEADING_BLACKLIST:
#             continue
#         # Multi-word match at the same position wins — it's more specific.
#         found.setdefault(m.start(), {"start": m.start(), "end": m.end(1) + 1, "heading": m.group(1)})

#     return [found[k] for k in sorted(found)]


# class DocxParser(BaseParser, Logger):
#     """Parser for DOCX with semantic chunking."""

#     # Paragraphs shorter than this word count are candidates for the
#     # structural-heading heuristic (regardless of .docx style).
#     _HEADING_MAX_WORDS: int = 8

#     # Chunks with fewer words than this are considered orphans and will be
#     # merged into a neighbouring chunk in the post-pass.
#     _ORPHAN_MIN_WORDS: int = 5

#     def __init__(self) -> None:
#         """Initialize the parser."""
#         super().__init__()
#         self.settings = get_settings()
#         from src.dependencies import get_service_container

#         self.service_container = get_service_container()
#         self.embedding_service: BaseEmbeddingService = self.service_container.embedding_service
#         self.vector_store = get_faiss_vector_store(self.embedding_service.get_embedding_dimensions())

#     @staticmethod
#     def _is_structural_heading(text: str, max_words: int = 8) -> bool:
#         """Return True when *text* looks like a title or section header even
#         though the .docx author did not apply a Heading style.

#         ALL must pass:
#             1. Word count <= max_words   — real headings are short.
#             2. Matches _SECTION_LABEL_RE — all-caps title".

#         """
#         words = text.split()
#         if len(words) > max_words or len(words) == 0:
#             return False
#         return bool(_SECTION_LABEL_RE.match(text.strip()))

#     @staticmethod
#     def cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
#         """Returns the similarity score between two vectors."""

#         return dot(vec1, vec2) / (norm(vec1) * norm(vec2))

#     @staticmethod
#     def _merge_orphan_chunks(chunks: List[str], min_words: int = 5) -> List[str]:
#         """Merge any chunk that is too short to stand alone into its neighbour."""

#         if not chunks:
#             return chunks

#         merged: List[str] = []
#         carry: str = ""

#         for chunk in chunks:
#             # Prepend any carried-over fragment
#             if carry:
#                 chunk = carry + " " + chunk
#                 carry = ""

#             if len(chunk.split()) < min_words:
#                 # Too short to emit — carry forward
#                 carry = chunk
#             else:
#                 merged.append(chunk)

#         # If something is still carried it means the very last chunk was short; append it to the previous chunk.
#         if carry:
#             if merged:
#                 merged[-1] = merged[-1] + " " + carry
#             else:
#                 merged.append(carry)

#         return merged

#     # Clause numbering patterns (e.g. "1.1 ", "2.3.4 ", "(a) ", "b) ")
#     _CLAUSE_PREFIX_RE = re.compile(r"^(\d+[\.\)]\d*[\.\d]*\s|" r"\([a-z]+\)\s|" r"[a-z]\)\s)")

#     def _clean_text(self, text: str) -> str:
#         """Clean the text to remove unwanted characters."""

#         if not text or not text.strip():
#             raise EmptyTextException("Text cannot be empty. Please provide valid text content.")

#         # Replace special whitespace chars
#         text = text.replace("\u00a0", " ")
#         text = text.replace("\u200b", "")
#         text = text.replace("\ufeff", "")
#         text = text.replace("\r", "")
#         text = re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]", "", text)

#         # Normalize whitespace
#         text = re.sub(r"\s+", " ", text).strip()

#         # Strip leading dots but preserve clause prefixes (e.g. "1.", "(a)")
#         stripped = text.lstrip(" \n\t")
#         if self._CLAUSE_PREFIX_RE.match(stripped):
#             text = stripped
#         else:
#             text = stripped.lstrip(".")

#         return text

#     async def clean_document(self, document: Document) -> None:
#         """Clean the document by removing trailing spaces and extra chars."""

#         try:
#             for paragraph in document.paragraphs:
#                 if paragraph.text:
#                     paragraph.text = " ".join(paragraph.text.split())
#         except Exception as e:
#             raise DocxCleaningException(str(e)) from e

#     async def _extract_metadata(self, document: Document) -> Dict[str, Any]:
#         """Extract document metadata."""

#         try:
#             self.logger.info("Extracting metadata from the document.")
#             properties = document.core_properties
#             paragraph_words = sum(len(p.text.split()) for p in document.paragraphs)
#             table_words = sum(len(cell.text.split()) for table in document.tables for row in table.rows for cell in row.cells)

#             return {
#                 "source": "docx",
#                 "author": properties.author or "Unknown",
#                 "title": properties.title or "Untitled",
#                 "created_at": properties.created.isoformat() if properties.created else None,
#                 "modified_at": properties.modified.isoformat() if properties.modified else None,
#                 "paragraph_count": len(document.paragraphs),
#                 "table_count": len(document.tables),
#                 "word_count": paragraph_words + table_words,
#             }

#         except Exception as e:
#             raise DocxMetadataExtractionException(str(e)) from e

#     async def _extract_tables(self, document: Document) -> List[Dict[str, Any]]:
#         try:
#             tables = []
#             for t_idx, table in enumerate(document.tables):
#                 self.logger.info("Cleaned text content for chunking.")
#                 rows = [[self._clean_text(cell.text) for cell in row.cells] for row in table.rows]
#                 tables.append(
#                     {
#                         "table_index": t_idx,
#                         "content": rows,
#                     }
#                 )
#             return tables
#         except Exception as e:
#             raise DocxTableExtractionException(str(e)) from e

#     async def _extract_paragraphs(self, document: Document) -> List[Dict[str, Any]]:
#         try:
#             data = []
#             for idx, p in enumerate(document.paragraphs):
#                 if p.text.strip():
#                     cleaned = self._clean_text(p.text)
#                     if cleaned:
#                         is_heading = bool((p.style and p.style.name.startswith("Heading")) or self._is_structural_heading(cleaned, self._HEADING_MAX_WORDS))
#                         data.append({"index": idx, "content": cleaned, "is_heading": is_heading})
#             return data
#         except Exception as e:
#             raise DocxParagraphExtractionException(str(e)) from e

#     @staticmethod
#     def _split_at_clause_boundaries(paragraphs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
#         """Split paragraphs at every inline clause heading.

#         Detects both multi-word titles (e.g. "Audit Rights.",
#         "Integration Testing.") and single-word titles (e.g. "Term.",
#         "Assignment.", "Definitions.") and splits at every boundary so that
#         each clause later becomes its own chunk.
#         """
#         result: List[Dict[str, Any]] = []

#         for para in paragraphs:
#             if para["is_heading"]:
#                 result.append(para)
#                 continue

#             text = para["content"]
#             boundaries = _find_clause_heading_matches(text)

#             if not boundaries:
#                 result.append(para)
#                 continue

#             segments: List[Dict[str, Any]] = []

#             # Text before the first clause heading belongs to the previous
#             # clause; emit as a plain (non-heading) paragraph.
#             if boundaries[0]["start"] > 0:
#                 prefix = text[: boundaries[0]["start"]].strip()
#                 if prefix:
#                     segments.append({"heading": None, "body": prefix})

#             # Each clause heading starts a segment that runs to the next
#             # heading (or end of paragraph).
#             for i, b in enumerate(boundaries):
#                 body_start = b["end"]
#                 body_end = boundaries[i + 1]["start"] if i + 1 < len(boundaries) else len(text)
#                 body = text[body_start:body_end].strip()
#                 segments.append({"heading": b["heading"], "body": body})

#             for seg in segments:
#                 if seg["heading"] is None:
#                     # Prefix of a mid-paragraph heading — belongs to previous clause.
#                     result.append(
#                         {
#                             "index": para["index"],
#                             "content": seg["body"],
#                             "is_heading": False,
#                         }
#                     )
#                     continue

#                 result.append(
#                     {
#                         "index": para["index"],
#                         "content": seg["heading"],
#                         "is_heading": True,
#                     }
#                 )
#                 if seg["body"]:
#                     result.append(
#                         {
#                             "index": para["index"],
#                             "content": seg["body"],
#                             "is_heading": False,
#                             "clause_boundary": True,
#                         }
#                     )

#         return result

#     async def _semantic_chunk_paragraphs(self, paragraphs: List[Dict[str, Any]]) -> List[str]:
#         """Do the semantic chunking for the paragraphs to be context aware."""

#         # Pre-process: split inline clause headings
#         paragraphs = self._split_at_clause_boundaries(paragraphs)

#         texts = [para["content"] for para in paragraphs]
#         embeddings = [await self.embedding_service.generate_embeddings(text=t, task="text-matching") for t in texts]

#         # Compute pairwise similarities between consecutive paragraphs
#         similarities = [self.cosine_similarity(embeddings[i], embeddings[i + 1]) for i in range(len(embeddings) - 1)]

#         # Split at points below threshold (mean - 0.75 * std)
#         mean_sim = np.mean(similarities)
#         std_sim = np.std(similarities)
#         threshold = mean_sim - 0.75 * std_sim
#         split_points = {i for i, sim in enumerate(similarities) if sim < threshold}

#         # Build chunks respecting headings, size limits, and split points
#         chunks: List[str] = []
#         chunk_headings: List[Optional[str]] = []
#         current: List[str] = []
#         current_len: int = 0
#         current_heading: Optional[str] = None
#         max_len: int = self.settings.chunk_size
#         pending_heading: Optional[str] = None

#         def _flush() -> None:
#             nonlocal current, current_len, pending_heading
#             if current:
#                 chunks.append(" ".join(current))
#                 chunk_headings.append(pending_heading)
#                 current = []
#                 current_len = 0

#         for i, para in enumerate(paragraphs):
#             text = para["content"]
#             text_len = len(text)

#             if para["is_heading"]:
#                 current_heading = text  # ← set FIRST
#                 _flush()  # ← then flush (pending_heading picks up new value inside)
#                 pending_heading = current_heading  # ← redundant but explicit

#             # rest of loop unchanged...
#             elif current_len + text_len > max_len:
#                 _flush()
#                 pending_heading = current_heading

#             if not current:
#                 pending_heading = current_heading

#             if not para["is_heading"]:  # ← don't append heading text as content
#                 current.append(text)
#                 current_len += text_len

#             if i in split_points or para.get("clause_boundary"):
#                 _flush()

#         _flush()

#         # Merge orphan chunks
#         raw_texts = self._merge_orphan_chunks(chunks, min_words=self._ORPHAN_MIN_WORDS)

#         # Re-align headings after merging
#         merged_headings: List[Optional[str]] = []
#         src_idx = 0
#         for merged_text in raw_texts:
#             if src_idx < len(chunk_headings):
#                 merged_headings.append(chunk_headings[src_idx])
#             else:
#                 merged_headings.append(None)

#             consumed = 0
#             for j in range(src_idx, len(chunks)):
#                 if chunks[j] in merged_text:
#                     consumed += 1
#                 else:
#                     break
#             src_idx += max(consumed, 1)

#         return [{"text": text, "section_heading": heading} for text, heading in zip(raw_texts, merged_headings)]

#     async def parse_data(self, data: List[TextInfo], session_data: Optional[Any] = None) -> ParseResult:
#         """Parse data using the registry services."""
#         start = time.time()

#         paragraphs = []
#         for d in data:
#             paragraphs.append(d.text)
#         if not paragraphs:
#             raise ValueError("No paragraphs found in the data.")

#         # Embed each paragraph and index it
#         for para in paragraphs:
#             text = para
#             if not text:
#                 continue
#             vector = await self.embedding_service.generate_embeddings(text=text, task="text-matching")
#             await self.vector_store.index_embedding(vector)

#         chunks: List[Chunk] = []

#         for i in range(len(paragraphs)):
#             chunk = Chunk(
#                 chunk_id=str(uuid.uuid4()),
#                 document_id=None,
#                 chunk_index=i,
#                 content=paragraphs[i],
#                 embedding_model=self.embedding_service.model_name,
#                 embedding_vector=vector,
#                 metadata={"chunk_type": "semantic_paragraph"},
#                 created_at=datetime.utcnow().isoformat(),
#             )
#             chunks.append(chunk)

#         return ParseResult(
#             success=True,
#             chunks=chunks,
#             metadata={"paragraph_count": len(paragraphs), "source": "parsed_data"},
#             processing_time=time.time() - start,
#         )

#     async def parse_document(self, document: Document, session_data: Optional["SessionData"] = None) -> ParseResult:
#         start = time.time()

#         try:
#             await self.clean_document(document)
#             metadata = await self._extract_metadata(document)

#             document_id = str(uuid.uuid4())
#             metadata["document_id"] = document_id

#             paragraphs = await self._extract_paragraphs(document)
#             tables = await self._extract_tables(document)

#             vector_store = session_data.vector_store if session_data else self.vector_store
#             semantic_chunks = await self._semantic_chunk_paragraphs(paragraphs)

#             chunks: List[Chunk] = []
#             chunk_index = 0

#             # Text chunks
#             for chunk_info in semantic_chunks:
#                 cleaned = self._clean_text(chunk_info["text"])
#                 if not cleaned:
#                     continue

#                 vector = await self.embedding_service.generate_embeddings(text=cleaned, task="text-matching")
#                 await vector_store.index_embedding(vector)

#                 chunk_metadata: Dict[str, Any] = {"chunk_type": "semantic_paragraph"}
#                 if chunk_info.get("section_heading"):
#                     chunk_metadata["section_heading"] = chunk_info["section_heading"]

#                 chunks.append(
#                     Chunk(
#                         chunk_id=str(uuid.uuid4()),
#                         document_id=document_id,
#                         chunk_index=chunk_index,
#                         content=cleaned,
#                         embedding_model=self.embedding_service.model_name,
#                         embedding_vector=None,
#                         metadata=chunk_metadata,
#                         created_at=datetime.utcnow().isoformat(),
#                     )
#                 )
#                 chunk_index += 1

#             # Table chunks
#             for table in tables:
#                 rows = [" | ".join(r) for r in table["content"]]
#                 table_text = self._clean_text(" ".join(rows))
#                 if not table_text:
#                     continue

#                 vector = await self.embedding_service.generate_embeddings(text=table_text)
#                 await vector_store.index_embedding(vector)

#                 chunks.append(
#                     Chunk(
#                         chunk_id=str(uuid.uuid4()),
#                         document_id=document_id,
#                         chunk_index=chunk_index,
#                         content=table_text,
#                         embedding_model=self.embedding_service.model_name,
#                         embedding_vector=None,
#                         metadata={
#                             "chunk_type": "table",
#                             "table_index": table["table_index"],
#                             "row_count": len(table["content"]),
#                         },
#                         created_at=datetime.utcnow().isoformat(),
#                     )
#                 )
#                 chunk_index += 1

#             # NOTE: chunk_store indexing is handled by IngestionService._parse_data()

#             return ParseResult(
#                 success=True,
#                 chunks=chunks,
#                 metadata=metadata,
#                 error_message=None,
#                 processing_time=time.time() - start,
#             )

#         except Exception as e:
#             self.logger.error(str(e))
#             return ParseResult(
#                 success=False,
#                 chunks=[],
#                 metadata={},
#                 error_message=str(e),
#                 processing_time=0.0,
#             )

#     def is_healthy(self) -> Any:
#         """Get the health status of the parser."""
#         pass

import re
import time
import uuid
from datetime import datetime
from typing import Any, Dict, List, Optional

import numpy as np
from docx.document import Document
from docx.text.paragraph import Paragraph
from numpy import dot
from numpy.linalg import norm

from src.config.logging import Logger
from src.config.settings import get_settings
from src.exceptions.parser_exceptions import (
    DocxCleaningException,
    DocxMetadataExtractionException,
    DocxParagraphExtractionException,
    DocxTableExtractionException,
    EmptyTextException,
)
from src.schemas.playbook_review import TextInfo
from src.schemas.registry import Chunk, ParseResult
from src.services.registry.base_parser import BaseParser
from src.services.session_manager import SessionData
from src.services.vector_store.embeddings.base_embedding_service import (
    BaseEmbeddingService,
)
from src.services.vector_store.manager import get_faiss_vector_store

# ==========================================================
# INLINE CLAUSE HEADING DETECTION
# ==========================================================

# Regex: multi-word inline clause headings like "Audit Rights. ", "Consulting Services.", "Non-Compete.", "Term/Termination."
_INLINE_CLAUSE_HEADING_RE = re.compile(
    r"""
    (?:(?:^|(?<=\.\s)|(?<=:\s)|(?<=;\s)))
    (
        [A-Z][A-Za-z0-9'\/-]+          # first capitalized word
        (?:\s+[A-Za-z0-9'\/-]+){0,6}   # allow lowercase connector words and shorter titles
    )
    \.\s
    """,
    re.VERBOSE,
)

# Regex: numbered inline clause headings like "2. Independent Company:" or "3A. Payments:"
_NUMBERED_INLINE_CLAUSE_HEADING_RE = re.compile(
    r"(?<!\S)"
    r"("  # heading capture
    r"\d+[A-Za-z]?"  # 1, 3A, 23a
    r"(?:\.\d+)*"  # 1.1, 2.3.4
    r"\.?\s*"  # optional dot + optional whitespace
    r"[A-Z][A-Za-z0-9'\-/ ]{1,120}?"  # title text
    r":"
    r")\s*"
)

# Regex: single-word clause titles like "Term. ", "Assignment. ", "Definitions. ".
# _SINGLE_WORD_HEADING_RE = re.compile(r"(?:^|(?<=\. ))" r"([A-Z][a-z]{3,}(?:'[a-z]+)?\.)" r"\s(?=[A-Z])")
_SINGLE_WORD_HEADING_RE = re.compile(r"(?:^|(?<=\.\s)|(?<=:\s)|(?<=;\s))" r"([A-Z][a-z]{3,}(?:'[a-z]+)?\.)" r"\s(?=[A-Z])")

# Words that look like single-word clause titles but are actually sentence adverbs/conjunctions
_SINGLE_WORD_HEADING_BLACKLIST = frozenset(
    {
        "however",
        "further",
        "furthermore",
        "additionally",
        "moreover",
        "otherwise",
        "nevertheless",
        "notwithstanding",
        "accordingly",
        "therefore",
        "thus",
        "hence",
        "thereby",
        "thereafter",
        "whereas",
        "specifically",
        "particularly",
        "collectively",
        "individually",
        "alternatively",
        "consequently",
        "subsequently",
        "respectively",
        "conversely",
        "similarly",
        "meanwhile",
        "indeed",
        "finally",
        "firstly",
        "secondly",
        "thirdly",
        "lastly",
        "instead",
    }
)

# Regex: all-caps structural headings only. Numeric clause headings are detected separately.
_SECTION_LABEL_RE = re.compile(r"^[A-Z][A-Z\s\.\,\&\'\-]{1,60}$")

_NUMBERED_CLAUSE_HEADING_RE = re.compile(
    r"""
    ^
    (
        \d+[A-Za-z]?            # 1, 3A, 23a
        (?:\.\d+)*              # 1.1, 2.3.4
    )
    \.?\s*                      # optional dot + optional space
    ([A-Z][^:]{1,100})          # Clause title
    :
    $
    """,
    re.VERBOSE,
)


def _find_clause_heading_matches(text: str) -> List[Dict[str, Any]]:
    """Return clause-heading matches in paragraph order, deduped by start position.

    Each entry is {"start": int, "end": int, "heading": str} where end points
    to the first character AFTER the heading (i.e. start of body text).
    """
    found: Dict[int, Dict[str, Any]] = {}

    for m in _NUMBERED_INLINE_CLAUSE_HEADING_RE.finditer(text):
        found[m.start(1)] = {
            "start": m.start(1),
            "end": m.end(1),
            "heading": m.group(1).rstrip(":"),
        }

    for m in _INLINE_CLAUSE_HEADING_RE.finditer(text):
        found[m.start(1)] = {
            "start": m.start(1),
            "end": m.end(),
            "heading": m.group(1),
        }

    for m in _SINGLE_WORD_HEADING_RE.finditer(text):
        word = m.group(1).rstrip(".").lower()
        if word in _SINGLE_WORD_HEADING_BLACKLIST:
            continue
        # Multi-word match at the same position wins — it's more specific.
        found.setdefault(
            m.start(1),
            {
                "start": m.start(1),
                "end": m.end(),
                "heading": m.group(1).rstrip("."),
            },
        )

    return [found[k] for k in sorted(found)]


# ==========================================================
# PRODUCTION LEGAL HEADING DETECTOR
# ==========================================================


class LegalHeadingDetector:

    NUMBERED_RE = re.compile(
        r"""
        ^
        (?P<number>\d+(?:\.\d+)*)
        [\.\)]\s+
        (?P<title>.+?)
        (?=\.\s|$)
        """,
        re.VERBOSE,
    )

    ALL_CAPS_RE = re.compile(r"^[A-Z][A-Z\s\-\&\,]{3,}$")
    MAX_HEADING_WORDS = 14

    @classmethod
    def detect(cls, paragraph: Paragraph, next_text: Optional[str]) -> Dict[str, Any]:

        raw_text = paragraph.text.strip()
        if not raw_text:
            return cls._empty()

        score = 0
        number = None
        title = None
        body = ""

        # 1️⃣ Numbered clause
        match = cls.NUMBERED_RE.match(raw_text)
        if match:
            score += 4
            number = match.group("number")
            title = match.group("title").strip()
            body = raw_text[match.end() :].strip()

        # 2️⃣ Word heading style
        if paragraph.style and paragraph.style.name.lower().startswith("heading"):
            score += 3
            title = raw_text

        # 3️⃣ ALL CAPS section
        if cls.ALL_CAPS_RE.match(raw_text):
            score += 2
            title = raw_text

        # 4️⃣ Mostly bold
        if cls._is_mostly_bold(paragraph):
            score += 1
            title = raw_text

        # 5️⃣ Structural short-line rule
        if cls._looks_structural(raw_text, next_text):
            score += 1
            title = raw_text

        is_heading = score >= 3
        confidence = min(score / 6, 1.0)

        return {
            "is_heading": is_heading,
            "confidence": confidence,
            "number": number,
            "title": title if is_heading else None,
            "body": body,
        }

    @staticmethod
    def _empty():
        return {
            "is_heading": False,
            "confidence": 0.0,
            "number": None,
            "title": None,
            "body": None,
        }

    @staticmethod
    def _is_mostly_bold(paragraph: Paragraph) -> bool:
        runs = paragraph.runs
        if not runs:
            return False

        bold_chars = sum(len(run.text) for run in runs if run.bold)
        total_chars = sum(len(run.text) for run in runs)
        return total_chars > 0 and (bold_chars / total_chars) > 0.6

    @classmethod
    def _looks_structural(cls, text: str, next_text: Optional[str]) -> bool:
        words = text.split()
        if len(words) == 0 or len(words) > cls.MAX_HEADING_WORDS:
            return False

        if text.count(".") > 1:
            return False

        if next_text and len(next_text.split()) > 20:
            return True

        if text.istitle() or text.isupper():
            return True

        return False


# ==========================================================
# DOCX PARSER
# ==========================================================


class DocxParser(BaseParser, Logger):

    _HEADING_MAX_WORDS: int = 8
    _ORPHAN_MIN_WORDS: int = 5

    def __init__(self) -> None:
        super().__init__()
        self.settings = get_settings()
        from src.dependencies import get_service_container

        self.service_container = get_service_container()
        self.embedding_service: BaseEmbeddingService = self.service_container.embedding_service
        self.vector_store = get_faiss_vector_store(self.embedding_service.get_embedding_dimensions())

    # --------------------------
    # Utilities
    # --------------------------

    @staticmethod
    def _is_structural_heading(text: str, max_words: int = 8) -> bool:
        """Return True when *text* looks like a title or section header even
        though the .docx author did not apply a Heading style.

        ALL must pass:
            1. Word count <= max_words   — real headings are short.
            2. Matches _SECTION_LABEL_RE — all-caps title".

        """
        words = text.split()
        if len(words) > max_words or len(words) == 0:
            return False
        return bool(_SECTION_LABEL_RE.match(text.strip()))

    @staticmethod
    def _looks_like_inline_heading(text: str) -> bool:
        """Treat short, title-cased legal clause lines ending in a period as headings."""
        stripped = text.strip()
        words = stripped.split()

        if len(words) > 8 or not stripped.endswith("."):
            return False

        lowercase_words = {"shall", "will", "agrees", "must", "may", "should", "can", "cannot"}
        if any(w in stripped.lower() for w in lowercase_words):
            return False

        return stripped[0].isupper()

    @staticmethod
    def cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
        return dot(vec1, vec2) / (norm(vec1) * norm(vec2))

    @staticmethod
    def cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
        return dot(vec1, vec2) / (norm(vec1) * norm(vec2))

    def _clean_text(self, text: str) -> str:
        if not text or not text.strip():
            raise EmptyTextException("Text cannot be empty.")

        text = text.replace("\u00a0", " ").replace("\u200b", "").replace("\ufeff", "")
        text = text.replace("\r", "")
        text = re.sub(r"[\x00-\x1f\x7f]", "", text)
        text = re.sub(r"\s+", " ", text).strip()
        return text

    async def clean_document(self, document: Document) -> None:
        try:
            for paragraph in document.paragraphs:
                if paragraph.text:
                    paragraph.text = " ".join(paragraph.text.split())
        except Exception as e:
            raise DocxCleaningException(str(e)) from e

    # --------------------------
    # Metadata
    # --------------------------

    async def _extract_metadata(self, document: Document) -> Dict[str, Any]:
        try:
            properties = document.core_properties
            return {
                "source": "docx",
                "author": properties.author or "Unknown",
                "title": properties.title or "Untitled",
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
            }
        except Exception as e:
            raise DocxMetadataExtractionException(str(e)) from e

    # --------------------------
    # Paragraph extraction (UPDATED FOR NUMBERED CLAUSES)
    # --------------------------

    async def _extract_paragraphs(self, document: Document) -> List[Dict[str, Any]]:
        try:
            data = []

            for idx, p in enumerate(document.paragraphs):
                raw = p.text.strip()
                if not raw:
                    continue

                cleaned = self._clean_text(raw)
                if not cleaned:
                    continue

                is_style_heading = bool(p.style and p.style.name and p.style.name.startswith("Heading"))

                is_structural = self._is_structural_heading(cleaned, self._HEADING_MAX_WORDS)

                is_numbered_clause = bool(_NUMBERED_CLAUSE_HEADING_RE.match(cleaned))

                is_inline_heading = self._looks_like_inline_heading(cleaned)

                is_heading = is_style_heading or is_structural or is_numbered_clause or is_inline_heading

                data.append(
                    {
                        "index": idx,
                        "content": cleaned,
                        "is_heading": is_heading,
                    }
                )

            return data

        except Exception as e:
            raise DocxParagraphExtractionException(str(e)) from e

    # --------------------------
    # Semantic Chunking (Clause-Aware)
    # --------------------------

    def _split_at_clause_boundaries(self, paragraphs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Split paragraphs at every inline clause heading.

        Detects both multi-word titles (e.g. "Audit Rights.",
        "Integration Testing.") and single-word titles (e.g. "Term.",
        "Assignment.", "Definitions.") and splits at every boundary so that
        each clause later becomes its own chunk.
        """
        result: List[Dict[str, Any]] = []

        for para in paragraphs:
            if para["is_heading"]:
                result.append(para)
                continue

            text = para["content"]
            boundaries = _find_clause_heading_matches(text)

            if not boundaries:
                result.append(para)
                continue

            segments: List[Dict[str, Any]] = []

            # Text before the first clause heading belongs to the previous clause
            if boundaries[0]["start"] > 0:
                prefix = text[: boundaries[0]["start"]].strip()
                if prefix:
                    segments.append({"heading": None, "body": prefix})

            # Each clause heading starts a segment that runs to the next heading
            for i, b in enumerate(boundaries):
                body_start = b["end"]
                body_end = boundaries[i + 1]["start"] if i + 1 < len(boundaries) else len(text)
                body = text[body_start:body_end].strip()
                segments.append({"heading": b["heading"], "body": body})

            for seg in segments:
                if seg["heading"] is None:
                    # Prefix of a mid-paragraph heading — belongs to previous clause
                    result.append(
                        {
                            "index": para["index"],
                            "content": seg["body"],
                            "is_heading": False,
                        }
                    )
                    continue

                result.append(
                    {
                        "index": para["index"],
                        "content": seg["heading"],
                        "is_heading": True,
                    }
                )
                if seg["body"]:
                    result.append(
                        {
                            "index": para["index"],
                            "content": seg["body"],
                            "is_heading": False,
                            "clause_boundary": True,
                        }
                    )

        return result

    async def _semantic_chunk_paragraphs(self, paragraphs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Clause-aware chunking.

        Every detected heading (style-based or inline clause heading)
        becomes the current section heading for subsequent paragraphs.
        Each non-heading paragraph becomes its own chunk.
        """

        paragraphs = self._split_at_clause_boundaries(paragraphs)

        chunks: List[Dict[str, Any]] = []
        current_heading: Optional[str] = None

        for para in paragraphs:
            text = para["content"].strip()
            if not text:
                continue

            if para["is_heading"]:
                current_heading = text
                continue

            chunks.append(
                {
                    "text": text,
                    "section_heading": current_heading,
                }
            )

        return chunks

    # --------------------------
    # Main Parse
    # --------------------------

    async def parse_document(
        self,
        document: Document,
        session_data: Optional["SessionData"] = None,
    ) -> ParseResult:

        start = time.time()

        try:
            await self.clean_document(document)
            metadata = await self._extract_metadata(document)

            document_id = str(uuid.uuid4())
            metadata["document_id"] = document_id

            paragraphs = await self._extract_paragraphs(document)
            semantic_chunks = await self._semantic_chunk_paragraphs(paragraphs)

            vector_store = session_data.vector_store if session_data else self.vector_store

            chunks: List[Chunk] = []
            chunk_index = 0

            for chunk_info in semantic_chunks:
                cleaned = self._clean_text(chunk_info["text"])
                if not cleaned:
                    continue

                vector = await self.embedding_service.generate_embeddings(
                    text=cleaned,
                    task="text-matching",
                )

                await vector_store.index_embedding(vector)

                metadata_chunk = {
                    "chunk_type": "semantic_paragraph",
                    "section_heading": chunk_info.get("section_heading"),
                }

                chunks.append(
                    Chunk(
                        chunk_id=str(uuid.uuid4()),
                        document_id=document_id,
                        chunk_index=chunk_index,
                        content=cleaned,
                        embedding_model=self.embedding_service.model_name,
                        embedding_vector=None,
                        metadata=metadata_chunk,
                        created_at=datetime.utcnow().isoformat(),
                    )
                )

                chunk_index += 1

            return ParseResult(
                success=True,
                chunks=chunks,
                metadata=metadata,
                error_message=None,
                processing_time=time.time() - start,
            )

        except Exception as e:
            self.logger.error(str(e))
            return ParseResult(
                success=False,
                chunks=[],
                metadata={},
                error_message=str(e),
                processing_time=0.0,
            )

    async def parse_data(self, data: List[TextInfo], session_data: Optional[Any] = None) -> ParseResult:
        """Parse data (list of TextInfo) without a document object."""
        start = time.time()

        try:
            if not data:
                raise ValueError("No data provided to parse.")

            vector_store = session_data.vector_store if session_data else self.vector_store
            chunks: List[Chunk] = []
            chunk_index = 0

            for text_info in data:
                text = text_info.text.strip() if hasattr(text_info, "text") else str(text_info).strip()
                if not text:
                    continue

                cleaned = self._clean_text(text)
                vector = await self.embedding_service.generate_embeddings(
                    text=cleaned,
                    task="text-matching",
                )
                await vector_store.index_embedding(vector)

                chunks.append(
                    Chunk(
                        chunk_id=str(uuid.uuid4()),
                        document_id=None,
                        chunk_index=chunk_index,
                        content=cleaned,
                        embedding_model=self.embedding_service.model_name,
                        embedding_vector=None,
                        metadata={"chunk_type": "parsed_text"},
                        created_at=datetime.utcnow().isoformat(),
                    )
                )
                chunk_index += 1

            return ParseResult(
                success=True,
                chunks=chunks,
                metadata={"text_count": len(data), "source": "parsed_data"},
                error_message=None,
                processing_time=time.time() - start,
            )

        except Exception as e:
            self.logger.error(str(e))
            return ParseResult(
                success=False,
                chunks=[],
                metadata={},
                error_message=str(e),
                processing_time=0.0,
            )

    def is_healthy(self) -> Any:
        return True
